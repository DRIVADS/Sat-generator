# ===================== IMPORTS =====================
import requests
import ssl
import datetime
import random
import string
from io import BytesIO


import qrcode
from barcode import Code128
from barcode.writer import ImageWriter

from bs4 import BeautifulSoup
from urllib3.poolmanager import PoolManager
from requests.adapters import HTTPAdapter

from docx import Document
from docx.shared import Inches


# ===================== SSL ADAPTER SAT =====================
class TLSAdapter(HTTPAdapter):
    def init_poolmanager(self, connections, maxsize, block=False):
        ctx = ssl.create_default_context()
        ctx.set_ciphers("DEFAULT@SECLEVEL=1")
        self.poolmanager = PoolManager(
            num_pools=connections,
            maxsize=maxsize,
            block=block,
            ssl_context=ctx
        )


# ===================== GENERADORES =====================
def generar_qr(url):
    qr = qrcode.make(url)
    buffer = BytesIO()
    qr.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def generar_codigo_barras_rfc(rfc):
    buffer = BytesIO()
    codigo = Code128(rfc, writer=ImageWriter())
    codigo.write(
        buffer,
        options={
            "module_width": 0.3,
            "module_height": 9,
            "quiet_zone": 6,
            "write_text": False
        }
    )
    buffer.seek(0)
    return buffer


def generar_cadena(rfc):
    fecha_hora = datetime.datetime.now().strftime("%Y/%m/%d %H:%M:%S")
    numero_largo = ''.join(random.choices(string.digits, k=21))
    base64_fake = ''.join(
        random.choices(string.ascii_letters + string.digits + '+/=', k=90)
    )
    return f"||{fecha_hora}|{rfc}|CONSTANCIA DE SITUACIÓN FISCAL|{numero_largo}|{base64_fake}||"


def generar_bloque_similar(longitud_total=180):
    caracteres = string.ascii_letters + string.digits + "+/="
    return ''.join(random.choices(caracteres, k=longitud_total))


# ===================== CONSULTA SAT =====================
def consultar_sat(id_cif, rfc):
    session = requests.Session()
    session.mount("https://", TLSAdapter())

    url = (
        "https://siat.sat.gob.mx/app/qr/faces/pages/mobile/"
        f"validadorqr.jsf?D1=10&D2=1&D3={id_cif}_{rfc}"
    )

    response = session.get(url, timeout=10)
    response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")

    datos = {}
    for fila in soup.select("tbody tr"):
        celdas = fila.find_all("td")
        if len(celdas) == 2:
            clave = celdas[0].get_text(strip=True).replace(":", "")
            valor = celdas[1].get_text(strip=True)
            datos[clave] = valor

    return datos


# ===================== UTILIDADES WORD =====================
def copiar_formato(destino, origen):
    destino.font.name = origen.font.name
    destino.font.size = origen.font.size
    destino.bold = origen.bold
    destino.italic = origen.italic
    destino.underline = origen.underline
    if origen.font.color.rgb:
        destino.font.color.rgb = origen.font.color.rgb


def reemplazar_placeholder_paragraph(paragraph, placeholder, valor, bold=None):
    texto_completo = "".join(run.text for run in paragraph.runs)

    if placeholder not in texto_completo:
        return

    base_run = paragraph.runs[0]
    paragraph.clear()

    partes = texto_completo.split(placeholder)

    run_antes = paragraph.add_run(partes[0])
    copiar_formato(run_antes, base_run)

    run_valor = paragraph.add_run(valor)
    copiar_formato(run_valor, base_run)

    if bold is not None:
        run_valor.bold = bold

    if len(partes) > 1:
        run_despues = paragraph.add_run(partes[1])
        copiar_formato(run_despues, base_run)


# ===================== FUNCIÓN PRINCIPAL =====================
def generar_constancia(
    plantilla,
    salida,
    id_cif,
    rfc,
    tamaño_qr=1.55,
    ancho_barras=2.07,
    alto_barras=0.39
):
    datos = consultar_sat(id_cif, rfc)

    meses = [
        "ENERO","FEBRERO","MARZO","ABRIL","MAYO","JUNIO",
        "JULIO","AGOSTO","SEPTIEMBRE","OCTUBRE","NOVIEMBRE","DICIEMBRE"
    ]
    hoy = datetime.date.today()
    fecha_formateada = f"{hoy.day} DE {meses[hoy.month-1]} DE {hoy.year}"

    key1 = generar_cadena(rfc)
    key2 = generar_bloque_similar()

    qr1 = generar_qr(
        f"https://siat.sat.gob.mx/app/qr/faces/pages/mobile/"
        f"validadorqr.jsf?D1=10&D2=1&D3={id_cif}_{rfc}"
    )

    qr2 = generar_qr(
        f"https://siat.sat.gob.mx/app/qr/faces/pages/mobile/"
        f"validadorqr.jsf?D1=26&D2=1&D3={random.randint(100000000,999999999)}_{rfc}"
    )

    barras = generar_codigo_barras_rfc(rfc)

    doc = Document(plantilla)

    reemplazos = {
        "{{RFC}}": rfc,
        "{{idCIF}}": id_cif,
        "{{DIAMESAÑO}}": fecha_formateada,
        "{{NOMBRE}}": datos.get("Nombre", ""),
        "{{APELLIDOPA}}": datos.get("Apellido Paterno", ""),
        "{{APELLIDOMA}}": datos.get("Apellido Materno", ""),
        "{{CURP}}": datos.get("CURP", ""),
        "{{FECHAINICIO}}": datos.get("Fecha de Inicio de operaciones","").replace("-","/"),
        "{{ESTATUSCONTRIBUYENTE}}": datos.get("Situación del contribuyente",""),
        "{{FECHACAMBIO}}": datos.get("Fecha del último cambio de situación","").replace("-","/"),
        "{{ENTIDADFEDERATIVA}}": datos.get("Entidad Federativa",""),
        "{{LOCALIDAD}}": datos.get("Localidad"," "),
        "{{MUNICIPIO}}": datos.get("Municipio o delegación",""),
        "{{COLONIA}}": datos.get("Colonia",""),
        "{{TIPOVIALIDAD}}": datos.get("Tipo de vialidad",""),
        "{{NOMBREVIALIDAD}}": datos.get("Nombre de la vialidad",""),
        "{{NUMEXTERIOR}}": datos.get("Número exterior",""),
        "{{NUMINTERIOR}}": datos.get("Número interior",""),
        "{{CODIGOPOSTAL}}": datos.get("CP",""),
        "{{CORREO}}": datos.get("Correo electrónico",""),
        "{{REGIMENFISCAL}}": datos.get("Régimen",""),
        "{{CALLECRUZE1}}": "",
        "{{FECHAALTA}}": datos.get("Fecha de alta","").replace("-","/"),
        "{{KEY1}}": key1,
        "{{KEY2}}": key2
    }

    sin_negrita = {
        "{{CALLECRUZE1}}","{{ENTIDADFEDERATIVA}}","{{LOCALIDAD}}",
        "{{MUNICIPIO}}","{{COLONIA}}","{{TIPOVIALIDAD}}",
        "{{NOMBREVIALIDAD}}","{{NUMEXTERIOR}}","{{NUMINTERIOR}}",
        "{{CODIGOPOSTAL}}"
    }

    for contenedor in [doc.paragraphs] + [
        cell.paragraphs
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    ]:
        for p in contenedor:
            for k, v in reemplazos.items():
                reemplazar_placeholder_paragraph(
                    p, k, v, bold=False if k in sin_negrita else None
                )

            if "[INSERTAR_QR_AQUI]" in p.text:
                p.clear()
                p.add_run().add_picture(qr1, Inches(tamaño_qr))

            if "[INSERTAR_validador_AQUI]" in p.text:
                p.clear()
                p.add_run().add_picture(qr2, Inches(tamaño_qr))

            if "[INSERTAR_BARRAS_RFC]" in p.text:
                p.clear()
                p.add_run().add_picture(barras, Inches(ancho_barras), Inches(alto_barras))

    doc.save(salida)
    print("✅ Documento generado correctamente")


# ===================== EJECUCIÓN =====================
# ===================== EJECUCIÓN =====================
if __name__ == "__main__":
    print("=== GENERADOR DE CONSTANCIA SAT ===")

    id_cif = input("Ingresa el ID CIF: ").strip()
    rfc = input("Ingresa el RFC: ").strip().upper()

    if not id_cif or not rfc:
        print("❌ El ID CIF y el RFC son obligatorios")
    else:
        generar_constancia(
            plantilla="plantilla.docx",
            salida="documento_final.docx",
            id_cif=id_cif,
            rfc=rfc
        )

#===============================================================